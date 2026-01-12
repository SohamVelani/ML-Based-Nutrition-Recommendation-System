import re
import json
import streamlit as st
from typing import Dict, List, Any, Optional, Tuple
import fitz  # PyMuPDF - better PDF processing
import docx
from io import BytesIO
import spacy
from spacy.matcher import Matcher
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import numpy as np
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class MedicalReportProcessor:
    def __init__(self):
        """Initialize the medical report processor with NLP model"""
        try:
            # Try to load a medical NLP model first, fallback to general English
            try:
                # Uncomment these lines if you have scispacy installed:
                # self.nlp = spacy.load("en_ner_bc5cdr_md")  # Medical NER model
                # logger.info("Loaded medical NER model (scispacy)")
                
                # For now, use the general English model
                self.nlp = spacy.load("en_core_web_sm")
                logger.info("Loaded general English NLP model")
            except OSError:
                try:
                    self.nlp = spacy.load("en_core_web_sm")
                    logger.info("Loaded fallback English NLP model")
                except OSError:
                    st.error("""
                    SpaCy English model not found! Please install it:
                    ```
                    python -m spacy download en_core_web_sm
                    ```
                    
                    For better medical text processing, also consider:
                    ```
                    pip install scispacy
                    pip install https://s3-us-west-2.amazonaws.com/ai2-s2-scispacy/releases/v0.5.1/en_ner_bc5cdr_md-0.5.1.tar.gz
                    ```
                    """)
                    self.nlp = None
                    return
            
            self.matcher = Matcher(self.nlp.vocab)
            self._setup_medical_patterns()
            self._setup_medical_ontology()
            
        except Exception as e:
            logger.error(f"Error initializing NLP model: {e}")
            self.nlp = None
    
    def _setup_medical_ontology(self):
        """Setup medical ontology and synonym mapping"""
        self.medical_ontology = {
            # Disease synonyms and abbreviations
            'diabetes': [
                'diabetes mellitus', 'dm', 't1dm', 't2dm', 'type 1 diabetes', 
                'type 2 diabetes', 'diabetic', 'insulin dependent diabetes',
                'non-insulin dependent diabetes', 'niddm', 'iddm'
            ],
            'hypertension': [
                'high blood pressure', 'htn', 'elevated blood pressure',
                'arterial hypertension', 'systolic hypertension', 'diastolic hypertension'
            ],
            'hyperlipidemia': [
                'high cholesterol', 'dyslipidemia', 'elevated cholesterol',
                'hypercholesterolemia', 'high ldl', 'low hdl'
            ],
            'obesity': [
                'overweight', 'increased bmi', 'elevated bmi', 'weight gain',
                'adiposity', 'morbid obesity'
            ],
            'cardiovascular_disease': [
                'heart disease', 'cardiac disease', 'coronary artery disease',
                'cad', 'myocardial infarction', 'heart attack', 'mi',
                'coronary heart disease', 'chd', 'ischemic heart disease'
            ],
            'kidney_disease': [
                'renal disease', 'chronic kidney disease', 'ckd', 'renal failure',
                'kidney failure', 'nephropathy', 'renal insufficiency',
                'end stage renal disease', 'esrd'
            ],
            'liver_disease': [
                'hepatic disease', 'liver dysfunction', 'cirrhosis',
                'fatty liver', 'nafld', 'hepatitis', 'liver failure'
            ],
            'thyroid_disorder': [
                'thyroid disease', 'hypothyroidism', 'hyperthyroidism',
                'thyroid dysfunction', 'goiter', 'thyroiditis'
            ]
        }
        
        # Negation patterns
        self.negation_patterns = [
            'no', 'not', 'without', 'absent', 'negative', 'denies',
            'rules out', 'ruled out', 'unlikely', 'no evidence of',
            'no signs of', 'no history of', 'unremarkable'
        ]
    
    def _setup_medical_patterns(self):
        """Setup enhanced patterns for medical information extraction"""
        
        # Enhanced medical conditions patterns with abbreviations
        conditions_patterns = [
            # Diabetes patterns
            [{"LOWER": {"IN": ["diabetes", "diabetic", "dm", "t1dm", "t2dm", "iddm", "niddm"]}},
             {"LOWER": {"IN": ["mellitus", "type", "1", "2", "i", "ii"]}, "OP": "*"}],
            
            # Hypertension patterns
            [{"LOWER": {"IN": ["hypertension", "htn"]}},
             {"LOWER": {"IN": ["stage", "grade"]}, "OP": "?"}, 
             {"LOWER": {"IN": ["1", "2", "i", "ii", "mild", "moderate", "severe"]}, "OP": "?"}],
            [{"LOWER": {"IN": ["high", "elevated"]}}, {"LOWER": {"IN": ["blood", "bp"]}}, {"LOWER": "pressure"}],
            
            # Other conditions with abbreviations
            [{"LOWER": {"IN": ["hyperlipidemia", "dyslipidemia", "hypercholesterolemia"]}}],
            [{"LOWER": {"IN": ["high", "elevated"]}}, {"LOWER": {"IN": ["cholesterol", "ldl"]}}],
            [{"LOWER": {"IN": ["obesity", "overweight", "increased", "elevated"]}}, {"LOWER": "bmi", "OP": "?"}],
            [{"LOWER": {"IN": ["cardiovascular", "cardiac", "coronary", "heart"]}}, 
             {"LOWER": {"IN": ["disease", "artery", "attack"]}}],
            [{"LOWER": {"IN": ["cad", "chd", "mi", "myocardial"]}}, {"LOWER": "infarction", "OP": "?"}],
            [{"LOWER": {"IN": ["kidney", "renal", "ckd"]}}, 
             {"LOWER": {"IN": ["disease", "failure", "dysfunction", "insufficiency"]}}],
            [{"LOWER": {"IN": ["liver", "hepatic", "nafld"]}}, 
             {"LOWER": {"IN": ["disease", "cirrhosis", "dysfunction", "failure"]}}],
            [{"LOWER": {"IN": ["thyroid", "hypothyroidism", "hyperthyroidism"]}},
             {"LOWER": {"IN": ["disorder", "disease", "dysfunction"]}, "OP": "?"}],
            [{"LOWER": {"IN": ["asthma", "copd", "emphysema", "bronchitis"]}}],
            [{"LOWER": "anemia"}],
            [{"LOWER": {"IN": ["depression", "anxiety", "bipolar", "ptsd"]}}],
            [{"LOWER": {"IN": ["arthritis", "osteoarthritis", "rheumatoid"]}}],
        ]
        
        # Enhanced lab values patterns
        lab_patterns = [
            # HbA1c patterns
            [{"LOWER": {"IN": ["hba1c", "a1c", "hemoglobin"]}}, {"LOWER": "a1c", "OP": "?"}, 
             {"IS_PUNCT": True, "OP": "?"}, {"LIKE_NUM": True}, {"TEXT": "%", "OP": "?"}],
            
            # Glucose patterns
            [{"LOWER": {"IN": ["glucose", "sugar", "fbs", "rbs", "fasting", "random"]}}, 
             {"LOWER": {"IN": ["glucose", "sugar"]}, "OP": "?"}, {"IS_PUNCT": True, "OP": "?"}, 
             {"LIKE_NUM": True}, {"LOWER": {"IN": ["mg/dl", "mmol/l"]}, "OP": "?"}],
            
            # Cholesterol patterns
            [{"LOWER": {"IN": ["cholesterol", "ldl", "hdl", "triglycerides", "tg"]}}, 
             {"IS_PUNCT": True, "OP": "?"}, {"LIKE_NUM": True}, 
             {"LOWER": {"IN": ["mg/dl", "mmol/l"]}, "OP": "?"}],
            
            # Blood pressure patterns
            [{"LOWER": {"IN": ["bp", "blood"]}}, {"LOWER": "pressure", "OP": "?"}, 
             {"IS_PUNCT": True, "OP": "?"}, {"LIKE_NUM": True}, 
             {"TEXT": {"IN": ["/", "over"]}, "OP": "?"}, {"LIKE_NUM": True}],
            
            # Weight and BMI
            [{"LOWER": {"IN": ["weight", "wt"]}}, {"IS_PUNCT": True, "OP": "?"}, 
             {"LIKE_NUM": True}, {"LOWER": {"IN": ["kg", "lb", "lbs", "pounds"]}, "OP": "?"}],
            [{"LOWER": {"IN": ["bmi", "body", "mass"]}}, {"LOWER": {"IN": ["mass", "index"]}, "OP": "?"}, 
             {"IS_PUNCT": True, "OP": "?"}, {"LIKE_NUM": True}],
        ]
        
        # Medication patterns with dosages
        medication_patterns = [
            # Diabetes medications
            [{"LOWER": {"IN": ["metformin", "insulin", "glipizide", "glyburide", "sitagliptin"]}}, 
             {"LIKE_NUM": True, "OP": "?"}, {"LOWER": {"IN": ["mg", "units", "ml", "mcg"]}, "OP": "?"}],
            
            # Cholesterol medications
            [{"LOWER": {"IN": ["atorvastatin", "simvastatin", "rosuvastatin", "pravastatin"]}}, 
             {"LIKE_NUM": True, "OP": "?"}, {"LOWER": "mg", "OP": "?"}],
            
            # Blood pressure medications
            [{"LOWER": {"IN": ["lisinopril", "amlodipine", "losartan", "metoprolol", "enalapril"]}}, 
             {"LIKE_NUM": True, "OP": "?"}, {"LOWER": "mg", "OP": "?"}],
        ]
        
        # Add patterns to matcher
        self.matcher.add("MEDICAL_CONDITIONS", conditions_patterns)
        self.matcher.add("LAB_VALUES", lab_patterns)
        self.matcher.add("MEDICATIONS", medication_patterns)
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Enhanced PDF text extraction with OCR fallback"""
        try:
            # First try PyMuPDF for better text extraction
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            text = ""
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                page_text = page.get_text()
                
                # Check if page has sufficient text (not scanned/image)
                if len(page_text.strip()) < 50:  # Threshold for "too little text"
                    logger.info(f"Page {page_num + 1} appears to be scanned, using OCR")
                    try:
                        # Convert page to image and use OCR
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        img = Image.open(BytesIO(img_data))
                        
                        # Use Tesseract OCR
                        ocr_text = pytesseract.image_to_string(img, config='--psm 6')
                        text += ocr_text + "\n"
                    except Exception as e:
                        logger.warning(f"OCR failed for page {page_num + 1}: {e}")
                        text += page_text + "\n"
                else:
                    text += page_text + "\n"
            
            pdf_document.close()
            return text
            
        except Exception as e:
            logger.error(f"Error reading PDF with PyMuPDF: {e}")
            # Fallback to pdf2image + OCR for entire document
            try:
                st.info("Using OCR for entire document...")
                images = convert_from_bytes(file_content)
                text = ""
                for i, image in enumerate(images):
                    logger.info(f"Processing page {i + 1} with OCR")
                    page_text = pytesseract.image_to_string(image, config='--psm 6')
                    text += page_text + "\n"
                return text
            except Exception as ocr_error:
                st.error(f"Both PDF extraction and OCR failed: {ocr_error}")
                return ""
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Enhanced DOCX text extraction"""
        try:
            doc_file = BytesIO(file_content)
            doc = docx.Document(doc_file)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                        
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def extract_text_from_file(self, uploaded_file) -> str:
        """Extract text from uploaded file with better error handling"""
        file_extension = uploaded_file.name.lower().split('.')[-1]
        file_content = uploaded_file.getvalue()
        
        try:
            if file_extension == 'pdf':
                return self.extract_text_from_pdf(file_content)
            elif file_extension in ['docx', 'doc']:
                return self.extract_text_from_docx(file_content)
            elif file_extension == 'txt':
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        return file_content.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                st.error("Unable to decode text file. Please ensure it's in UTF-8 format.")
                return ""
            else:
                st.error("Unsupported file format. Please upload PDF, DOCX, or TXT files.")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_extension}: {e}")
            st.error(f"Error processing file: {str(e)}")
            return ""
    
    def _detect_negation(self, text: str, keyword: str, context_window: int = 10) -> bool:
        """Detect if a medical condition is negated"""
        words = text.lower().split()
        
        try:
            keyword_idx = words.index(keyword.lower())
            
            # Check words before the keyword within context window
            start_idx = max(0, keyword_idx - context_window)
            preceding_words = words[start_idx:keyword_idx]
            
            # Check for negation patterns
            for neg_pattern in self.negation_patterns:
                for word in preceding_words:
                    if neg_pattern in word or word in neg_pattern:
                        return True
            
            return False
        except ValueError:
            return False
    
    def _normalize_condition(self, condition_text: str) -> Tuple[str, str]:
        """Normalize medical condition using ontology"""
        condition_lower = condition_text.lower().strip()
        
        # Check against ontology
        for standard_name, synonyms in self.medical_ontology.items():
            if condition_lower in [s.lower() for s in synonyms]:
                return standard_name, condition_text
            
            # Partial matching for complex terms
            for synonym in synonyms:
                if synonym.lower() in condition_lower or condition_lower in synonym.lower():
                    return standard_name, condition_text
        
        return condition_text, condition_text
    
    def extract_medical_info(self, text: str) -> Dict[str, Any]:
        """Enhanced medical information extraction with negation detection"""
        if not self.nlp:
            return {"error": "NLP model not loaded"}
        
        # Process text with spaCy
        doc = self.nlp(text)
        
        # Find matches using patterns
        matches = self.matcher(doc)
        
        extracted_info = {
            'medical_conditions': [],
            'lab_values': [],
            'medications': [],
            'vital_signs': [],
            'allergies': [],
            'summary_keywords': [],
            'raw_matches': [],
            'negated_conditions': []  # Track negated conditions
        }
        
        # Process pattern matches
        for match_id, start, end in matches:
            match_label = self.nlp.vocab.strings[match_id]
            matched_text = doc[start:end].text
            context = self._extract_full_sentence_context(text, matched_text)
            
            extracted_info['raw_matches'].append({
                'type': match_label,
                'text': matched_text,
                'context': context,
                'start': start,
                'end': end
            })
        
        # Enhanced condition extraction with negation detection
        text_lower = text.lower()
        
        # Extract conditions using ontology
        for standard_name, synonyms in self.medical_ontology.items():
            for synonym in synonyms:
                if synonym.lower() in text_lower:
                    # Check for negation
                    is_negated = self._detect_negation(text_lower, synonym)
                    
                    if is_negated:
                        extracted_info['negated_conditions'].append({
                            'condition': standard_name,
                            'original_text': synonym,
                            'context': self._extract_full_sentence_context(text_lower, synonym)
                        })
                    else:
                        # Check if already added
                        existing_conditions = [c['condition'] if isinstance(c, dict) else c 
                                             for c in extracted_info['medical_conditions']]
                        
                        if standard_name not in existing_conditions:
                            context = self._extract_full_sentence_context(text_lower, synonym)
                            extracted_info['medical_conditions'].append({
                                'condition': standard_name,
                                'original_text': synonym,
                                'context': context,
                                'confidence': 'high' if synonym.lower() == standard_name else 'medium'
                            })
        
        # Enhanced lab values extraction with better error handling
        lab_patterns = {
            'hba1c': r'(?:hba1c|a1c|hemoglobin\s*a1c)[\s:]*([0-9]+\.?[0-9]*)\s*%?',
            'glucose': r'(?:glucose|sugar|fbs|rbs|fasting\s*glucose|random\s*glucose)[\s:]*([0-9]+\.?[0-9]*)\s*(?:mg/dl|mmol/l)?',
            'total_cholesterol': r'(?:total\s*cholesterol|cholesterol)[\s:]*([0-9]+\.?[0-9]*)\s*(?:mg/dl|mmol/l)?',
            'ldl': r'(?:ldl|low\s*density\s*lipoprotein)[\s:]*([0-9]+\.?[0-9]*)\s*(?:mg/dl|mmol/l)?',
            'hdl': r'(?:hdl|high\s*density\s*lipoprotein)[\s:]*([0-9]+\.?[0-9]*)\s*(?:mg/dl|mmol/l)?',
            'triglycerides': r'(?:triglycerides|tg)[\s:]*([0-9]+\.?[0-9]*)\s*(?:mg/dl|mmol/l)?',
            'blood_pressure': r'(?:bp|blood\s*pressure)[\s:]*([0-9]+)\s*[/\\]\s*([0-9]+)',
            'weight': r'(?:weight|wt)[\s:]*([0-9]+\.?[0-9]*)\s*(?:kg|lb|lbs|pounds)?',
            'bmi': r'(?:bmi|body\s*mass\s*index)[\s:]*([0-9]+\.?[0-9]*)',
            'creatinine': r'(?:creatinine|cr)[\s:]*([0-9]+\.?[0-9]*)\s*(?:mg/dl|umol/l)?',
            'hemoglobin': r'(?:hemoglobin|hb|hgb)[\s:]*([0-9]+\.?[0-9]*)\s*(?:g/dl|g/l)?'
        }
        
        for lab_name, pattern in lab_patterns.items():
            matches = re.finditer(pattern, text_lower)
            for match in matches:
                try:
                    if lab_name == 'blood_pressure':
                        systolic = int(match.group(1))
                        diastolic = int(match.group(2))
                        value = f"{systolic}/{diastolic}"
                        
                        # Validate blood pressure ranges
                        if 60 <= systolic <= 250 and 40 <= diastolic <= 150:
                            extracted_info['lab_values'].append({
                                'test': lab_name,
                                'value': value,
                                'systolic': systolic,
                                'diastolic': diastolic,
                                'context': match.group(0),
                                'interpretation': self._interpret_blood_pressure(systolic, diastolic)
                            })
                    else:
                        value_str = match.group(1)
                        try:
                            value = float(value_str)
                            
                            # Basic validation based on typical ranges
                            if self._validate_lab_value(lab_name, value):
                                extracted_info['lab_values'].append({
                                    'test': lab_name,
                                    'value': value_str,
                                    'numeric_value': value,
                                    'context': match.group(0),
                                    'interpretation': self._interpret_lab_value(lab_name, value)
                                })
                        except ValueError:
                            logger.warning(f"Could not convert {value_str} to float for {lab_name}")
                            
                except Exception as e:
                    logger.warning(f"Error processing lab value {lab_name}: {e}")
                    continue
        
        # Enhanced medication extraction
        medication_keywords = [
            # Diabetes medications
            'metformin', 'insulin', 'glipizide', 'glyburide', 'sitagliptin', 'pioglitazone',
            # Hypertension medications
            'lisinopril', 'amlodipine', 'losartan', 'metoprolol', 'enalapril', 'hydrochlorothiazide',
            # Cholesterol medications
            'atorvastatin', 'simvastatin', 'rosuvastatin', 'pravastatin',
            # Other common medications
            'aspirin', 'warfarin', 'levothyroxine', 'omeprazole'
        ]
        
        for medication in medication_keywords:
            if medication in text_lower:
                context = self._extract_full_sentence_context(text_lower, medication)
                if context and not self._detect_negation(text_lower, medication):
                    # Extract dosage if present
                    dosage_pattern = rf'{medication}[\s,]*([0-9]+\.?[0-9]*)\s*(mg|mcg|units|ml)'
                    dosage_match = re.search(dosage_pattern, text_lower)
                    
                    med_info = {
                        'medication': medication,
                        'context': context
                    }
                    
                    if dosage_match:
                        med_info['dosage'] = dosage_match.group(1)
                        med_info['unit'] = dosage_match.group(2)
                    
                    extracted_info['medications'].append(med_info)
        
        # Generate enhanced summary keywords
        extracted_info['summary_keywords'] = self._generate_enhanced_summary_keywords(extracted_info)
        
        return extracted_info
    
    def _extract_full_sentence_context(self, text: str, keyword: str) -> str:
        """Extract full sentence containing the keyword"""
        sentences = re.split(r'[.!?]+', text)
        
        for sentence in sentences:
            if keyword.lower() in sentence.lower():
                return sentence.strip()
        
        # Fallback to word-based context if sentence splitting fails
        return self._extract_context(text, keyword, 50)
    
    def _validate_lab_value(self, lab_name: str, value: float) -> bool:
        """Validate lab values against typical ranges"""
        validation_ranges = {
            'hba1c': (3.0, 20.0),  # % 
            'glucose': (50, 600),   # mg/dl
            'total_cholesterol': (100, 500),  # mg/dl
            'ldl': (50, 400),       # mg/dl
            'hdl': (20, 150),       # mg/dl
            'triglycerides': (50, 1000),  # mg/dl
            'weight': (30, 300),    # kg
            'bmi': (10, 60),        # kg/m²
            'creatinine': (0.5, 15), # mg/dl
            'hemoglobin': (5, 20)   # g/dl
        }
        
        if lab_name in validation_ranges:
            min_val, max_val = validation_ranges[lab_name]
            return min_val <= value <= max_val
        
        return True  # If no range defined, accept the value
    
    def _interpret_lab_value(self, lab_name: str, value: float) -> str:
        """Provide interpretation of lab values"""
        interpretations = {
            'hba1c': {
                'ranges': [(0, 5.7, 'Normal'), (5.7, 6.5, 'Prediabetes'), (6.5, 20, 'Diabetes')],
                'unit': '%'
            },
            'glucose': {
                'ranges': [(0, 100, 'Normal'), (100, 126, 'Prediabetes'), (126, 600, 'Diabetes')],
                'unit': 'mg/dl (fasting)'
            },
            'total_cholesterol': {
                'ranges': [(0, 200, 'Desirable'), (200, 240, 'Borderline High'), (240, 500, 'High')],
                'unit': 'mg/dl'
            },
            'ldl': {
                'ranges': [(0, 100, 'Optimal'), (100, 130, 'Near Optimal'), (130, 160, 'Borderline High'), (160, 400, 'High')],
                'unit': 'mg/dl'
            },
            'bmi': {
                'ranges': [(0, 18.5, 'Underweight'), (18.5, 25, 'Normal'), (25, 30, 'Overweight'), (30, 60, 'Obese')],
                'unit': 'kg/m²'
            }
        }
        
        if lab_name in interpretations:
            ranges = interpretations[lab_name]['ranges']
            unit = interpretations[lab_name]['unit']
            
            for min_val, max_val, interpretation in ranges:
                if min_val <= value < max_val:
                    return f"{interpretation} ({unit})"
            
            return f"Outside normal range ({unit})"
        
        return "Normal range unknown"
    
    def _interpret_blood_pressure(self, systolic: int, diastolic: int) -> str:
        """Interpret blood pressure values"""
        if systolic < 120 and diastolic < 80:
            return "Normal"
        elif systolic < 130 and diastolic < 80:
            return "Elevated"
        elif (120 <= systolic <= 129) or (80 <= diastolic <= 89):
            return "Stage 1 Hypertension"
        elif systolic >= 130 or diastolic >= 90:
            return "Stage 2 Hypertension"
        else:
            return "Hypertensive Crisis"
    
    def _extract_context(self, text: str, keyword: str, context_length: int = 50) -> str:
        """Extract context around a keyword (fallback method)"""
        index = text.find(keyword.lower())
        if index == -1:
            return ""
        
        start = max(0, index - context_length)
        end = min(len(text), index + len(keyword) + context_length)
        context = text[start:end].strip()
        
        # Clean up context
        context = re.sub(r'\s+', ' ', context)
        return context
    
    def _generate_enhanced_summary_keywords(self, extracted_info: Dict[str, Any]) -> List[str]:
        """Generate enhanced summary keywords from extracted information"""
        keywords = []
        
        # Add conditions
        for condition in extracted_info['medical_conditions']:
            if isinstance(condition, dict):
                keywords.append(condition['condition'])
            else:
                keywords.append(condition)
        
        # Add significant lab values with interpretations
        for lab in extracted_info['lab_values']:
            test_name = lab['test']
            
            try:
                if test_name == 'hba1c':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 6.5:
                        keywords.append('diabetes_confirmed')
                    elif value >= 5.7:
                        keywords.append('prediabetes')
                        
                elif test_name == 'glucose':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 126:
                        keywords.append('elevated_glucose')
                        
                elif test_name == 'total_cholesterol':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 240:
                        keywords.append('high_cholesterol')
                        
                elif test_name == 'bmi':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 30:
                        keywords.append('obesity')
                        
                elif test_name == 'blood_pressure':
                    if 'systolic' in lab and 'diastolic' in lab:
                        systolic = lab['systolic']
                        diastolic = lab['diastolic']
                        if systolic >= 130 or diastolic >= 90:
                            keywords.append('hypertension')
                            
            except (ValueError, KeyError) as e:
                logger.warning(f"Error processing lab value {test_name}: {e}")
                continue
        
        return list(set(keywords))  # Remove duplicates
    
    def format_medical_summary(self, extracted_info: Dict[str, Any]) -> str:
        """Format extracted medical information into a comprehensive readable summary"""
        if 'error' in extracted_info:
            return f"Error: {extracted_info['error']}"
        
        summary = "📋 **Medical Report Summary**\n\n"
        
        # Medical Conditions
        if extracted_info['medical_conditions']:
            summary += "🏥 **Medical Conditions:**\n"
            for condition in extracted_info['medical_conditions']:
                if isinstance(condition, dict):
                    condition_name = condition['condition'].replace('_', ' ').title()
                    confidence = condition.get('confidence', 'medium')
                    original = condition.get('original_text', '')
                    
                    summary += f"• {condition_name}"
                    if confidence == 'high':
                        summary += " ✓"
                    elif confidence == 'medium':
                        summary += " ~"
                    
                    if original and original != condition['condition']:
                        summary += f" (found as: {original})"
                    summary += "\n"
                else:
                    summary += f"• {condition.replace('_', ' ').title()}\n"
            summary += "\n"
        
        # Negated Conditions (important for accuracy)
        if extracted_info.get('negated_conditions'):
            summary += "❌ **Ruled Out Conditions:**\n"
            for neg_condition in extracted_info['negated_conditions']:
                condition_name = neg_condition['condition'].replace('_', ' ').title()
                summary += f"• No {condition_name}\n"
            summary += "\n"
        
        # Lab Results with interpretations
        if extracted_info['lab_values']:
            summary += "🧪 **Lab Results:**\n"
            for lab in extracted_info['lab_values']:
                test_name = lab['test'].replace('_', ' ').upper()
                value = lab['value']
                interpretation = lab.get('interpretation', '')
                
                summary += f"• {test_name}: {value}"
                if interpretation:
                    summary += f" - {interpretation}"
                summary += "\n"
            summary += "\n"
        
        # Medications
        if extracted_info['medications']:
            summary += "💊 **Current Medications:**\n"
            for med in extracted_info['medications']:
                med_name = med['medication'].title()
                summary += f"• {med_name}"
                
                if 'dosage' in med and 'unit' in med:
                    summary += f" {med['dosage']} {med['unit']}"
                summary += "\n"
            summary += "\n"
        
        # Key Health Indicators
        if extracted_info['summary_keywords']:
            summary += "🔍 **Key Health Indicators:**\n"
            for keyword in extracted_info['summary_keywords']:
                formatted_keyword = keyword.replace('_', ' ').title()
                summary += f"• {formatted_keyword}\n"
        
        return summary
    
    def get_diet_relevant_conditions(self, extracted_info: Dict[str, Any]) -> List[str]:
        """Extract conditions that are relevant for diet recommendations with enhanced mapping"""
        diet_relevant_conditions = []
        
        # Enhanced mapping from medical conditions to diet considerations
        diet_mapping = {
            'diabetes': 'diabetes',
            'diabetes_confirmed': 'diabetes',
            'prediabetes': 'prediabetes',
            'hypertension': 'high_blood_pressure',
            'hyperlipidemia': 'high_cholesterol',
            'high_cholesterol': 'high_cholesterol',
            'obesity': 'weight_management',
            'cardiovascular_disease': 'heart_disease',
            'kidney_disease': 'kidney_disease',
            'liver_disease': 'liver_disease',
            'thyroid_disorder': 'thyroid_disorder'
        }
        
        # Process medical conditions
        for condition in extracted_info['medical_conditions']:
            condition_text = condition if isinstance(condition, str) else condition['condition']
            
            for medical_term, diet_term in diet_mapping.items():
                if medical_term in condition_text.lower():
                    if diet_term not in diet_relevant_conditions:
                        diet_relevant_conditions.append(diet_term)
        
        # Check lab values for diet-relevant indicators
        for lab in extracted_info['lab_values']:
            test_name = lab['test']
            
            try:
                if test_name == 'hba1c':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 6.5:
                        if 'diabetes' not in diet_relevant_conditions:
                            diet_relevant_conditions.append('diabetes')
                    elif value >= 5.7:
                        if 'prediabetes' not in diet_relevant_conditions:
                            diet_relevant_conditions.append('prediabetes')
                            
                elif test_name == 'glucose':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 126:
                        if 'diabetes' not in diet_relevant_conditions:
                            diet_relevant_conditions.append('diabetes')
                    elif value >= 100:
                        if 'prediabetes' not in diet_relevant_conditions:
                            diet_relevant_conditions.append('prediabetes')
                            
                elif test_name == 'total_cholesterol':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 200:
                        if 'high_cholesterol' not in diet_relevant_conditions:
                            diet_relevant_conditions.append('high_cholesterol')
                            
                elif test_name == 'ldl':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 130:
                        if 'high_cholesterol' not in diet_relevant_conditions:
                            diet_relevant_conditions.append('high_cholesterol')
                            
                elif test_name == 'bmi':
                    value = float(lab.get('numeric_value', lab['value']))
                    if value >= 30:
                        if 'weight_management' not in diet_relevant_conditions:
                            diet_relevant_conditions.append('weight_management')
                    elif value >= 25:
                        if 'weight_management' not in diet_relevant_conditions:
                            diet_relevant_conditions.append('mild_weight_management')
                            
                elif test_name == 'blood_pressure':
                    if 'systolic' in lab and 'diastolic' in lab:
                        systolic = lab['systolic']
                        diastolic = lab['diastolic']
                        if systolic >= 130 or diastolic >= 90:
                            if 'high_blood_pressure' not in diet_relevant_conditions:
                                diet_relevant_conditions.append('high_blood_pressure')
                                
            except (ValueError, KeyError) as e:
                logger.warning(f"Error processing lab value {test_name} for diet mapping: {e}")
                continue
        
        # Check summary keywords
        for keyword in extracted_info.get('summary_keywords', []):
            if keyword in diet_mapping:
                diet_term = diet_mapping[keyword]
                if diet_term not in diet_relevant_conditions:
                    diet_relevant_conditions.append(diet_term)
        
        return list(set(diet_relevant_conditions))  # Remove duplicates
    
    def get_detailed_health_profile(self, extracted_info: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a detailed health profile for more sophisticated diet recommendations"""
        health_profile = {
            'risk_factors': [],
            'metabolic_status': {},
            'cardiovascular_risk': 'unknown',
            'dietary_priorities': [],
            'restrictions': [],
            'monitoring_needed': []
        }
        
        # Analyze metabolic status
        for lab in extracted_info.get('lab_values', []):
            test_name = lab['test']
            
            try:
                if test_name == 'hba1c':
                    value = float(lab.get('numeric_value', lab['value']))
                    health_profile['metabolic_status']['hba1c'] = {
                        'value': value,
                        'status': 'diabetic' if value >= 6.5 else 'prediabetic' if value >= 5.7 else 'normal'
                    }
                    
                elif test_name == 'glucose':
                    value = float(lab.get('numeric_value', lab['value']))
                    health_profile['metabolic_status']['glucose'] = {
                        'value': value,
                        'status': 'diabetic' if value >= 126 else 'prediabetic' if value >= 100 else 'normal'
                    }
                    
                elif test_name in ['total_cholesterol', 'ldl', 'hdl']:
                    value = float(lab.get('numeric_value', lab['value']))
                    health_profile['metabolic_status'][test_name] = {
                        'value': value,
                        'interpretation': lab.get('interpretation', '')
                    }
                    
            except (ValueError, KeyError):
                continue
        
        # Determine cardiovascular risk
        risk_factors = 0
        conditions = [c if isinstance(c, str) else c['condition'] 
                     for c in extracted_info.get('medical_conditions', [])]
        
        if any('diabetes' in c or 'hypertension' in c for c in conditions):
            risk_factors += 2
        if any('cholesterol' in c for c in conditions):
            risk_factors += 1
        if any('obesity' in c for c in conditions):
            risk_factors += 1
            
        if risk_factors >= 3:
            health_profile['cardiovascular_risk'] = 'high'
        elif risk_factors >= 1:
            health_profile['cardiovascular_risk'] = 'moderate'
        else:
            health_profile['cardiovascular_risk'] = 'low'
        
        
        # Generate dietary priorities
        if 'diabetes' in conditions or health_profile['metabolic_status'].get('hba1c', {}).get('status') == 'diabetic':
            health_profile['dietary_priorities'].extend(['blood_sugar_control', 'carbohydrate_management'])
            
        if 'hypertension' in conditions:
            health_profile['dietary_priorities'].extend(['sodium_reduction', 'dash_diet'])
            
        if any('cholesterol' in c for c in conditions):
            health_profile['dietary_priorities'].extend(['saturated_fat_reduction', 'fiber_increase'])
            
        if 'obesity' in conditions:
            health_profile['dietary_priorities'].extend(['calorie_control', 'portion_management'])
        
        return health_profile